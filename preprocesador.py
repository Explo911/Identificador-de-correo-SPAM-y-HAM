import concurrent.futures
import os
from docx import Document
import xlsxwriter
import openpyxl
from nltk.stem.snowball import SnowballStemmer
import shutil
import threading
import math

# las variables path solo son para poder cambiar facilmente la ubicacion de los documentos resultantes
# en caso de tener que mover/renombrar carpetas
# HAM_folder = "KNN-ENTRENAMIENTO/HAM"
# SPAM_folder = "KNN-ENTRENAMIENTO/SPAM"
# path_resultados_HAM = "PreprocesamientoSteps/HAM"
# path_resultados_SPAM = "PreprocesamientoSteps/SPAM"
stopwords = open("stopwords.txt")
lista_stopwords = stopwords.read().split('\n')
stopwords.close()
stemmer = SnowballStemmer("spanish")


# preprocesamiento de la consulta
# realiza eliminacion de stopword y stemming
def minipreprocesador(consulta):
    temp = []
    lista_palabras = {}
    for word in consulta.split():
        word = word.lower()
        for letra in word:
            if letra.isalpha:
                if letra == "á":
                    temp.append("a")
                elif letra == "é":
                    temp.append("e")
                elif letra == "í":
                    temp.append("i")
                elif letra == "ó":
                    temp.append("o")
                elif letra == "ú" or letra == "ü":
                    temp.append("u")
                else:
                    temp.append(letra)
        palabra = ''.join(temp)
        if len(palabra) > 0:
            if palabra in lista_palabras:
                num = lista_palabras[palabra] + 1
            else:
                num = 1
            lista_palabras.update({palabra: num})
        temp = []
    temp = lista_palabras.copy()
    for palabra in temp.keys():
        if palabra in lista_stopwords:
            lista_palabras.pop(palabra)

    lista_consulta = {}
    for palabra, num in list(lista_palabras.items()):
        old_stem = palabra
        while True:
            new_stem = stemmer.stem(old_stem)
            if new_stem == old_stem:
                break
            else:
                old_stem = new_stem
        if old_stem in lista_consulta:
            cantidad = lista_consulta[old_stem] + num
        else:
            cantidad = num
        lista_consulta.update({old_stem: cantidad})
    print(lista_consulta)
    valores = list(lista_consulta.values())
    maximo = max(valores)
    new_valores = normal_min_max(valores, 0, maximo)
    i = 0
    for k in lista_consulta.keys():
        lista_consulta.update({k: new_valores[i]})
        i += 1
    return lista_consulta


# obtiene una lista de stems de todo el corpus
def stems_corpus(path_resultados):
    dataframe = openpyxl.load_workbook(f'{path_resultados}/4DiccStems.xlsx')
    df = dataframe.active
    stems = []
    for row in df.iter_rows(0, df.max_row):
        stems.append(row[0].value)
    return stems


# obtiene la matrix TF IDF
def matriz_TFIDF(lista_docs, path_resultados, normalize):
    stem_matrix = openpyxl.load_workbook(f"{path_resultados}/5MatrixBinariaDeStems.xlsx")
    sm = stem_matrix.active
    matriz = xlsxwriter.Workbook(f"{path_resultados}/MatrizTF-IDFNormStems.xlsx")
    matriz_worksheet = matriz.add_worksheet()
    stems = stems_corpus(path_resultados)

    # aqui calcula nk, el numero de documentos donde aparece cada termino
    nk = []
    columnas = list(sm.columns)[1:]
    i = 1
    for column in columnas:
        count = 0
        for cell in column:
            if cell.value == 1:
                count += 1
        nk.append(count)
        # aprovechamos para escribir cada uno de los stems en la primera fila
        matriz_worksheet.write(0, i, column[0].value)
        i += 1

    N = len(lista_docs)
    # obtengo valor maximo para realizar normalizacion min max
    dataframe = openpyxl.load_workbook(f"{path_resultados}/4DiccStems.xlsx")
    df = dataframe.active
    maximo = df['B1'].value
    dataframe.close()
    # esto es lo mas pesado de la funcion
    i = 1
    y = 1
    for doc in lista_docs:
        matriz_worksheet.write(i, 0, doc)
        i += 1
        # abrimos uno de los diccionarios especificos
        doc_list = openpyxl.load_workbook(f"{path_resultados}\{doc}")
        dl = doc_list.active
        # sacamos cada uno de los valores del diccionario y lo guardamos
        row_dicc = {}
        for row in dl.iter_rows(0, dl.max_row):
            row_dicc.update({row[0].value: row[1].value})
        # iteramos por cada uno de los stems de la lista completa
        fila = []
        j = 0
        for stem in stems:
            # checa si el stem existe en el diccionario
            if stem in row_dicc:
                # obtenemos el valor tf * idf
                tf = row_dicc[stem]
                idf = math.log(N / nk[j], 2)
                # print(stem, tf, tf*idf)
                fila.append(tf * idf)
            else:
                fila.append(0)
            j += 1
        # se normaliza la fila resultante
        if normalize:
            normal_fila = normal_min_max(fila, 0, maximo)
        else:
            normal_fila = fila
        x = 1
        # se escribe en el excel
        for norm in normal_fila:
            matriz_worksheet.write(y, x, norm)
            x += 1
        y += 1
    matriz.close()


def join_matrix(HAM, SPAM, result):
    H_dataframe = openpyxl.load_workbook(f"{HAM}/MatrizTF-IDFNormStems.xlsx")
    hdf = H_dataframe.active
    S_dataframe = openpyxl.load_workbook(f"{SPAM}/MatrizTF-IDFNormStems.xlsx")
    sdf = S_dataframe.active
    H_stems = stems_corpus(HAM)
    S_stems = stems_corpus(SPAM)
    total_stems = list(set(H_stems) | set(S_stems))
    matrix_total = xlsxwriter.Workbook(f"{result}/MatrizTF-IDF.xlsx")
    matrix_worksheet = matrix_total.add_worksheet()
    i = 1
    for stem in total_stems:
        matrix_worksheet.write(0, i, stem)
        i += 1
    j = 1
    matrix_dicc = []
    for row in hdf.iter_rows(2, hdf.max_row):
        row_dicc = {}
        cells = list(row)
        name = cells[0].value
        full_row = list(zip(H_stems, cells[1:]))
        for s, c in full_row:
            row_dicc.update({s: c.value})
        for stem in S_stems:
            row_dicc.update({stem: 0})
        matrix_worksheet.write(j, 0, name)
        j += 1
        matrix_dicc.append(row_dicc)
    for row in sdf.iter_rows(2, sdf.max_row):
        row_dicc = {}
        cells = list(row)
        name = cells[0].value
        full_row = list(zip(S_stems, cells[1:]))
        for s, c in full_row:
            row_dicc.update({s: c.value})
        for stem in H_stems:
            row_dicc.update({stem: 0})
        matrix_worksheet.write(j, 0, name)
        j += 1
        matrix_dicc.append(row_dicc)
    maximo = 0
    for dicc in matrix_dicc:
        if maximo < max(dicc.values()):
            maximo = max(dicc.values())
    i = 1
    for dicc in matrix_dicc:
        j = 1
        for stem in total_stems:
            normalized = (dicc[stem] - 0) / maximo
            matrix_worksheet.write(i, j, normalized)
            j += 1
        if i < 12:
            matrix_worksheet.write(i, j, "H")
        else:
            matrix_worksheet.write(i, j, "S")
        i += 1
    matrix_total.close()


# normaliza la matriz con metodo min-max
def normal_min_max(fila, min, max):
    new_fila = []
    for f in fila:
        new_fila.append((f - min) / (max - min))
    return new_fila


def extrae_tabla(diccionario, tablas):
    for tabla in tablas:
        for row in tabla.rows:
            if row.cells[0].tables:
                subtabla = []
                for cell in row.cells:
                    for t in cell.tables:
                        subtabla.append(t)
                extrae_tabla(diccionario, subtabla)
            texto = [cell.text for cell in row.cells]
            for parrafo in texto:
                parrafo = parrafo.lower()
                for palabra in parrafo.split():
                    temp = []
                    for letter in palabra:
                        if letter.isalpha() or letter == "@":
                            if letter == "á":
                                letter = 'a'
                            if letter == "é":
                                letter = 'e'
                            if letter == "í":
                                letter = 'i'
                            if letter == "ó":
                                letter = 'o'
                            if letter == "ú" or letter == "ü":
                                letter = 'u'
                            temp.append(letter)
                    word = ''.join(temp)
                    if len(word) > 0:
                        if word in diccionario:
                            num = diccionario[word] + 1
                        else:
                            num = 1
                        diccionario.update({word: num})


# tokeniza bien el documento
# se hace desde cero ignorando lo resultante del primer tokenizador
# solo porque me resultó más facil
def tokenizar_limpieza(folder_name, doc_name, path_resultados):
    print(f"-----{doc_name}-----")
    doc = Document(f"{folder_name}/{doc_name}")
    lista_palabras = {}
    workbook = xlsxwriter.Workbook(f"{path_resultados}/{doc_name.replace('.docx', '')}_2DiccMinus.xlsx")
    worksheet = workbook.add_worksheet()
    for par in doc.paragraphs:
        texto = par.text.split()
        for word in texto:
            palabra = []
            for letter in word:
                # solo dejamos pasar letras
                if letter.isalpha():
                    # nos deshacemos de letras con acento
                    letter = letter.lower()
                    if letter == "á":
                        letter = 'a'
                    if letter == "é":
                        letter = 'e'
                    if letter == "í":
                        letter = 'i'
                    if letter == "ó":
                        letter = 'o'
                    if letter == "ú" or letter == "ü":
                        letter = 'u'
                    palabra.append(letter)
                else:
                    break
            palabra = ''.join(palabra)
            if len(palabra) > 0:
                if palabra in lista_palabras:
                    num = lista_palabras[palabra] + 1
                else:
                    num = 1
                lista_palabras.update({palabra: num})

    tabla = doc.tables
    if tabla:
        print("Extrae info de tabla")
        extrae_tabla(lista_palabras, tabla)

    lista_final = sorted(lista_palabras.items(), key=lambda x: x[1], reverse=True)
    row = 0
    for elemento in lista_final:
        p, n = elemento
        worksheet.write(row, 0, p)
        worksheet.write(row, 1, n)
        row += 1
    workbook.close()


# eliminar stopwords
def no_stopwords(doc_name, path_resultados):
    dataframe = openpyxl.load_workbook(f'{path_resultados}\{doc_name.replace(".docx", "")}_2DiccMinus.xlsx')
    df = dataframe.active
    workbook = xlsxwriter.Workbook(f"{path_resultados}\{doc_name.replace('.docx', '')}_3DiccSinStopWords.xlsx")
    worksheet = workbook.add_worksheet()
    n = 0
    for row in df.iter_rows(0, df.max_row):
        # comparamos con archivo .txt de stopwords
        if not (row[0].value in lista_stopwords):
            worksheet.write(n, 0, row[0].value)
            worksheet.write(n, 1, row[1].value)
            n += 1
    workbook.close()


# realizamos stemming
def stemming(doc_name, path_resultados):
    path = f'{path_resultados}\{doc_name.replace(".docx", "")}'
    dataframe = openpyxl.load_workbook(f'{path}_3DiccSinStopWords.xlsx')
    df = dataframe.active
    workbook = xlsxwriter.Workbook(f'{path}_4DiccSteams.xlsx')
    worksheet = workbook.add_worksheet()
    fila = 0
    lista_stems = {}
    for row in df.iter_rows(0, df.max_row):
        new_stem = ""
        old_stem = row[0].value
        # este ciclo while nos asegura que consigamos stem correcto
        while True:
            new_stem = stemmer.stem(old_stem)
            # si es igual significa que no se puede reducir más
            if new_stem == old_stem:
                break
            else:
                old_stem = new_stem
        if new_stem in lista_stems:
            m = lista_stems[new_stem]
            lista_stems.update({new_stem: m + row[1].value})
        else:
            lista_stems.update({new_stem: row[1].value})
    lista_final = sorted(lista_stems.items(), key=lambda x: x[1], reverse=True)
    for stem in lista_final:
        p, n = stem
        worksheet.write(fila, 0, p)
        worksheet.write(fila, 1, n)
        fila += 1
    workbook.close()


# junta todos los documentos individuales en un mega documento para todo el corpus
# funcion solo hace esto para cada categoria, i.e. 1Diccionario.xlsx
def unificar(doc_name, lista_docs, path_resultados):
    workbook = xlsxwriter.Workbook(f'{path_resultados}\{doc_name}')
    worksheet = workbook.add_worksheet()
    dicc = {}
    for doc in lista_docs:
        dataframe = openpyxl.load_workbook(f'{path_resultados}\{doc}')
        df = dataframe.active
        for row in df.iter_rows(0, df.max_row):
            if row[0].value in dicc:
                num = dicc[row[0].value] + row[1].value
            else:
                num = row[1].value
            dicc.update({row[0].value: num})
        dataframe.close()
    dicc_final = sorted(dicc.items(), key=lambda x: x[1], reverse=True)
    fila = 0
    for item in dicc_final:
        p, n = item
        worksheet.write(fila, 0, p)
        worksheet.write(fila, 1, n)
        fila += 1
    workbook.close()


# forma matriz binaria
# index_completo -> todos los indices individuales (4DiccStems.xlsx)
# index_individual -> mega indice de todo el corpus
def matriz_binaria(index_completo, index_individual, doc_name, path_resultados):
    matriz = xlsxwriter.Workbook(f'{path_resultados}\{doc_name}')
    worksheet = matriz.add_worksheet()

    complete = openpyxl.load_workbook(f'{path_resultados}\{index_completo}')
    cdf = complete.active
    lista_stems = []
    col = 1
    # se escribe todos los stems en la primera fila del excel
    for row in cdf.iter_rows(0, cdf.max_row):
        worksheet.write(0, col, row[0].value)
        lista_stems.append(row[0].value)
        col += 1

    i = 1
    # escribe cada fila de la matriz
    for ind in index_individual:
        index = openpyxl.load_workbook(f'{path_resultados}\{ind}')
        indexdf = index.active
        # escribe el nombre del documento en la primera casilla
        # el nombre lo escribe como "documento_4DiccStems.xlsx" por flojera mia
        worksheet.write(i, 0, ind)
        lista_temp = []
        # se obtienen todos los stems de dicho indice
        for row in indexdf.iter_rows(0, indexdf.max_row):
            lista_temp.append(row[0].value)
        col = 1
        # se recorre la lista de stems total y se compara con la lista de stems del indice especifico
        for stem in lista_stems:
            worksheet.write(i, col, 1 if stem in lista_temp else 0)
            col += 1
        index.close()
        i += 1
    matriz.close()


# lo junte todo en una misma funcion para poder intentar multithreading
def todo_el_proceso(folder_name, doc_name, path_resultados):
    tokenizar_limpieza(folder_name, doc_name, path_resultados)
    no_stopwords(doc_name, path_resultados)
    stemming(doc_name, path_resultados)


def main(doc_folder, resultado_folder, normalizar=False):
    lista_corpus = os.listdir(doc_folder)
    threads = concurrent.futures.ThreadPoolExecutor(max_workers=15)
    # mi intento de multithreading que no se si funcionó
    for doc in lista_corpus:
        threads.submit(todo_el_proceso(doc_folder, doc, resultado_folder))
    threads.shutdown(wait=True)

    lista_resultados = os.listdir(resultado_folder)
    # una forma extraña de guardar cada lista de documentos
    lista_docs = [[], [], [], []]
    # esta lista son los documentos que quiero que NO se eliminen
    lista_dics = ["1Diccionario.xlsx", "2DiccMinus.xlsx", "3DiccSinStopwords.xlsx", "4DiccStems.xlsx",
                  "5ListDiccIndex.xlsx", "5MatrixBinariaDeStems.xlsx", "MatrizTF-IDFNormStems.xlsx"]
    for res in lista_resultados:
        if "_1Dicc" in res:
            lista_docs[0].append(res)
        if "_2Dicc" in res:
            lista_docs[1].append(res)
        if "_3Dicc" in res:
            lista_docs[2].append(res)
        if "_4Dicc" in res:
            lista_docs[3].append(res)
    unificar(lista_dics[0], lista_docs[0], resultado_folder)
    unificar(lista_dics[1], lista_docs[1], resultado_folder)
    unificar(lista_dics[2], lista_docs[2], resultado_folder)
    unificar(lista_dics[3], lista_docs[3], resultado_folder)

    matriz_binaria(lista_dics[3], lista_docs[3], "5MatrixBinariaDeStems.xlsx", resultado_folder)
    matriz_TFIDF(lista_docs[3], resultado_folder, normalizar)
    # elimina todos los documentos intermediarios para solo dejar los mega indices
    for res in lista_resultados:
        if not (res in lista_dics):
            os.remove(f'{resultado_folder}\{res}')

    print("Listo")
